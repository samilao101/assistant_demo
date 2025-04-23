
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt


#HEADER
def create_headings(doc):
    doc.add_heading("CHANGES TO EXISTING SUPPLIER: GSAR EFORM STEPS ", level=1).style = doc.styles["Heading 1"]
    doc.add_heading("Find the supplier account: ", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("Go to: https://cummins-prods360int.mdm.informaticahosted.com/mdmapps/mdm/entity360view/")
    doc.add_paragraph("Click on GSAR Supplier")
    doc.add_picture('GSAR_functions/eform_images/heading1.JPG', width=Pt(450), height=Pt(225))  # Adjust width as needed
    doc.add_paragraph("Click on Queries > Supplier_SimpleSearch")
    doc.add_picture('GSAR_functions/eform_images/heading2.JPG', width=Pt(450), height=Pt(225))  # Adjust width as needed
    doc.add_paragraph("Enter the supplier number > Click on Run > Double Click on the supplier name")
    doc.add_picture('GSAR_functions/eform_images/heading3.JPG', width=Pt(450), height=Pt(225))  # Adjust width as needed
    doc.add_paragraph("Click on the + under GSAR Request Details > Enter Title> Change Request type to Changes to Existing Supplier > Enter the OnBehalf OF")
    doc.add_picture('GSAR_functions/eform_images/heading4.JPG', width=Pt(450), height=Pt(225))  # Adjust width as needed
    doc.add_paragraph("Scroll down and enter GSAR Operating Unit (where the supplier will be used) and the commodity or category")
    doc.add_picture('GSAR_functions/eform_images/heading5.JPG', width=Pt(450), height=Pt(225))  # Adjust width as needed

#US TAX ID
def create_ustax_id(doc):
    doc.add_heading("Confirm Tax Confirmation:", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("Under Supplier Header Tax, confirm the tax information matches. If there is not tax info add it. If the tax information has changed, you may need to create a new supplier (if the supplier is not in the system already).")
    doc.add_picture('GSAR_functions/eform_images/tax1.JPG', width=Pt(450), height=Pt(225))

#SIGNATURE
def create_signature(doc):
    doc.add_heading("Submit the GSAR eForm:", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("When you are ready to submit, remove the Draft checkmark, click Submit")
    doc.add_picture('GSAR_functions/eform_images/signature1.JPG', width=Pt(450), height=Pt(225))


#CHANGES:

#NAME CHANGE:
def name_update(doc):
    doc.add_heading("Name Update:", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("Update the Supplier Name to match the tax documentation")
    doc.add_picture('GSAR_functions/eform_images/namechange1.JPG', width=Pt(450), height=Pt(225))
    doc.add_paragraph("Answer question Is there a supplier name change?")
    doc.add_picture('GSAR_functions/eform_images/namechange2.JPG', width=Pt(450), height=Pt(225))
    doc.add_paragraph("Under Supplier Header Attachments, attach the letterhead with the name and change relevant tax document.")
    doc.add_picture('GSAR_functions/eform_images/attachments1.JPG', width=Pt(450), height=Pt(225))

#ADDRESS CHANGE:
def address_replace_update(doc):
    doc.add_heading("Address Update:", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("Scroll to Supplier Header Address > find the address being updated > click on the pencil > update it > click on the checkmark to save")
    doc.add_picture('GSAR_functions/eform_images/addresschange1.JPG', width=Pt(450), height=Pt(225))
    doc.add_paragraph("Under Supplier Header Attachments, attach the letterhead with the address change.")
    doc.add_picture('GSAR_functions/eform_images/attachments1.JPG', width=Pt(450), height=Pt(225))


#BANKING CHANGE
def banking_update(doc):
    doc.add_heading("Banking Update:", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("Under GSAR Request Details answer question Will the supplier be paid via EFT or Wire Transfer")
    doc.add_picture('GSAR_functions/eform_images/bankingupdate1.JPG', width=Pt(450), height=Pt(225))
    doc.add_paragraph("Under Supplier Header Attachments, attach the letterhead with the new banking. Make sure it has the attachment type of EFT/Wire....")
    doc.add_picture('GSAR_functions/eform_images/attachments1.JPG', width=Pt(450), height=Pt(225))

#CONTACT UPDATE

def contact_replace_update(doc):
    doc.add_heading("CHANGES TO EXISTING SUPPLIER", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("(The completed sections of the template must be affixed within the body of the supplier's company letterhead)")
    doc.add_picture('GSAR_functions/eform_images/logo.png')  # Adjust width as needed

 
#NEW CONTACT
def contact_add_additional(doc):
    doc.add_heading("CHANGES TO EXISTING SUPPLIER", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("(The completed sections of the template must be affixed within the body of the supplier's company letterhead)")
    doc.add_picture('GSAR_functions/eform_images/logo.png')  # Adjust width as needed


#ADD ADDRESS 
def address_add_additional(doc):
    doc.add_heading("CHANGES TO EXISTING SUPPLIER", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("(The completed sections of the template must be affixed within the body of the supplier's company letterhead)")
    doc.add_picture('GSAR_functions/eform_images/logo.png')  # Adjust width as needed

#CURRENCY UPDATE
def currency_update(doc):
    doc.add_heading("CHANGES TO EXISTING SUPPLIER", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("(The completed sections of the template must be affixed within the body of the supplier's company letterhead)")
    doc.add_picture('GSAR_functions/eform_images/logo.png')  # Adjust width as needed