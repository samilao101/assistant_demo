
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor


#HEADER
def create_headings(doc):
    doc.add_heading("CHANGES TO EXISTING SUPPLIER", level=1).style = doc.styles["Heading 1"]
    doc.add_paragraph("(The completed sections of the template must be affixed within the body of the supplier's company letterhead)")


#US TAX ID
def create_ustax_id(doc):
    tblNew = doc.add_table(rows=4, cols=1)
    tblNew.style = "Table Grid"

    headings = [
        "SUPPLIER NAME (as legally registered/appears on W-9 tax documentation:)",
        "",  # Leave a blank row
        "TAXPAYER ID (As provided on W-9 Tax Document)",
        ""
    ]

    for i, text in enumerate(headings):
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text

#SIGNATURE
def create_signature(doc):
    doc.add_paragraph()
    signature_heading = doc.add_paragraph("SIGNATURE ")
    signature_heading.style = doc.styles["Heading 2"]

    tblNew = doc.add_table(rows=8, cols=1)
    tblNew.style = "Table Grid"

    cell_texts = [
        "All pages of the supplier's letterhead and any additional supporting document(s) must be dated (within one year), contain the printed name, job title, eSignature or physical signature (not fancy font) of the supplier's company representative providing the information.",
        "",  # Leave a blank row
        "eSignature or Physical Signature:",
        "Printed Name:",
        "Job Title:",
        "Date:",
        "",
        "The completed sections of the template must be affixed within the body of the supplier's company letterhead"
    ]

    for i, text in enumerate(cell_texts):
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text


#CHANGES:

#NAME CHANGE:
def name_update(doc):
    doc.add_paragraph()
    type_of_change_heading = doc.add_paragraph("TYPE OF CHANGE:  SUPPLIER NAME CHANGE (change supplier name 'from', 'to')")
    type_of_change_heading.style = doc.styles["Heading 2"]

    tblNew = doc.add_table(rows=2, cols=2)  # Three rows: FROM, TO, and Additional Row
    tblNew.style = "Table Grid"

    cell_texts = [
        "FROM: PREVIOUS SUPPLIER NAME CURRENTLY LISTED IN THE CUMMINS SUPPLIER MASTER (To be completed by Cummins contact)",
        "TO: NEW SUPPLIER NAME" # Add your additional row text for column 2
    ]

    for i, text in enumerate(cell_texts):
        row_idx = i // 2  # Divide by 2 to determine row index
        col_idx = i % 2   # Use modulo 2 to determine column index
        cell = tblNew.cell(row_idx, col_idx)
        cell.paragraphs[0].text = text



#ADDRESS CHANGE:

def address_replace_update(doc):
    doc.add_paragraph()
    type_of_change_heading = doc.add_paragraph("TYPE OF CHANGE:  EXISTING ADDRESS CHANGE (change address 'from', 'to')")
    type_of_change_heading.style = doc.styles["Heading 2"]

    tblNew = doc.add_table(rows=15, cols=2)
    tblNew.style = "Table Grid"

    cell_texts_left = [
        "FROM: ADDRESS DETAILS CURRENTLY LISTED IN THE CUMMINS SUPPLIER MASTER (To be completed by Cummins contact).",
        "Address Line 1:",
        "Address Line 2:",
        "Address Line 3:",
        "City:",
        "County:",
        "Province:",
        "State:",
        "Postal Code:",
        "Country:",
        "",
        "ADDRESS PURPOSE: (Mark address purpose below)",
        "__ PURCHASING ADDRESS",
        "__ PAY ADDRESS",
        "Definition for 'PURCHASING ADDRESS': For Indirect Suppliers it is the PURCHASE ORDERS MAILING ADDRESS, For Direct suppliers it is the SHIP FROM ADDRESS."
    ]

    for i, text in enumerate(cell_texts_left):
 
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text


    cell_texts_right = [
        "TO: CHANGE WITH THE FOLLOWING INFORMATION",
        "Address Line 1:",
        "Address Line 2:",
        "Address Line 3:",
        "City:",
        "County:",
        "Province:",
        "State:",
        "Postal Code:",
        "Country:",
        "",
        "ADDRESS PURPOSE: (Mark address purpose below)",
        "__ PURCHASING ADDRESS",
        "__ PAY ADDRESS",
        "Definition for 'PURCHASING ADDRESS': For Indirect Suppliers it is the PURCHASE ORDERS MAILING ADDRESS, For Direct suppliers it is the SHIP FROM ADDRESS."
    ]

    for i, text in enumerate(cell_texts_right):
 
        cell = tblNew.cell(i, 1)
        cell.paragraphs[0].text = text



#BANKING CHANGE
def banking_update(doc):
    doc.add_paragraph()
    type_of_change_heading = doc.add_paragraph("TYPE OF CHANGE: BANKING INFORMATION")
    type_of_change_heading.style = doc.styles["Heading 2"]

    tblNew = doc.add_table(rows=17, cols=1)
    tblNew.style = "Table Grid"

    cell_texts = [
        "NEW BANKING DETAILS",
        "Bank Name:",
        "Beneficiary Name:",
        "Routing Number (ABA) (US based suppliers only):",
        "Bank Account Number:",
        "Beneficiary A/C No./ IBAN No.:",
        "Swift Code:",
        "Transit No (Canada based suppliers only):",
        "Institution No (Canada based suppliers only):",
        "Sort Code (UK based suppliers only):",
        "ACH Code (if applicable):",
        "Bank City:",
        "Bank State:",
        "Bank Country:",
        "Bank Agency Number (if applicable):",
        "Currency:",
        "EMAIL ADDRESS OF PERSON TO RECEIVE REMITTANCE ADVICE NOTIFICATION (when electronic payments are issued):"
    ]

    for i, text in enumerate(cell_texts):
 
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text



#CONTACT UPDATE

def contact_replace_update(doc):
    doc.add_paragraph()
    type_of_change_heading = doc.add_paragraph("TYPE OF CHANGE: CONTACT INFORMATION CHANGE ")
    type_of_change_heading.style = doc.styles["Heading 2"]

    tblNew = doc.add_table(rows=34, cols=2)
    tblNew.style = "Table Grid"

    cell_texts_left = [
        "PREVIOUS PURCHASE ORDER CONTACT CURRENTLY LISTED IN THE CUMMINS SUPPLIER MASTER: (To be completed by Cummins contact)",
        "First/Last Name:",
        "Email Address:",
        "Phone Number:",
        "Fax Number (If applicable):",
        "",
        "NEW/ADDITIONAL PURCHASE ORDER CONTACT DETAILS",
        "First/Last Name:",
        "Email Address:",
        "Phone Number:",
        "Fax Number (If applicable):",
        "",
        "",
        "",
        "",
        "",
        "PREVIOUS ACCOUNTS RECEIVABLE CONTACT (CURRENTLY LISTED IN THE CUMMINS SUPPLIER MASTER)",
        "First/Last Name:",
        "Email Address:",
        "Phone Number:",
        "Fax Number (If applicable):",
        "",
        "NEW/ADDITIONAL ACCOUNTS RECEIVABLE CONTACT DETAILS",
        "First/Last Name:",
        "Email Address:",
        "Phone Number:",
        "Fax Number (If applicable):",
        "",
        "",
        "",
        "",
        "",
        "",
        ""
    ]

    for i, text in enumerate(cell_texts_left):
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text

    cell_texts_right  = [
        "",
        "",
        "",
        "",
        "",
        "",
        "PURCHASING ADDRESS THIS CONTACT SHOULD BE ASSOCIATED WITH",
        "Address Line 1:",
        "Address Line 2:",
        "Address Line 3:",
        "City:",
        "County:",
        "Province:",
        "State:",
        "Postal Code:",
        "Country:",
        "",
        "",
        "",
        "",
        "",
        "",
        "PAY ADDRESS ADDRESS THIS CONTACT SHOULD BE ASSOCIATED WITH",
        "Address Line 1:",
        "Address Line 2:",
        "Address Line 3:",
        "City:",
        "County:",
        "Province:",
        "State:",
        "Postal Code:",
        "Country:"
    ]

    for i, text in enumerate(cell_texts_right):
        cell = tblNew.cell(i, 1)
        cell.paragraphs[0].text = text

 
#NEW CONTACT
def contact_add_additional(doc):
    doc.add_paragraph()
    type_of_change_heading = doc.add_paragraph("TYPE OF CHANGE: CONTACT INFORMATION ADDITION")
    type_of_change_heading.style = doc.styles["Heading 2"]

    tblNew = doc.add_table(rows=21, cols=2)
    tblNew.style = "Table Grid"

    cell_texts_left = [
    "NEW/ADDITIONAL PURCHASE ORDER CONTACT DETAILS",
    "First/Last Name:",
    "Email Address:",
    "Phone Number:",
    "Fax Number (If applicable):",
    "",
    "",
    "",
    "",
    "",
    "",
    "NEW/ADDITIONAL ACCOUNTS RECEIVABLE CONTACT DETAILS",
    "First/Last Name:",
    "Email Address:",
    "Phone Number:",
    "Fax Number (If applicable):",
    "",
    "",
    "",
    "",
    ""
    ]


    for i, text in enumerate(cell_texts_left):
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text

    cell_texts_right  = [
    "PURCHASING ADDRESS THIS CONTACT SHOULD BE ASSOCIATED WITH",
    "Address Line 1:",
    "Address Line 2:",
    "Address Line 3:",
    "City:",
    "County:",
    "Province:",
    "State:",
    "Postal Code:",
    "Country:",
    "",
    "PAY ADDRESS THIS CONTACT SHOULD BE ASSOCIATED WITH",
    "Address Line 1:",
    "Address Line 2:",
    "Address Line 3:",
    "City:",
    "County:",
    "Province:",
    "State:",
    "Postal Code:",
    "Country:"
    ] 

    for i, text in enumerate(cell_texts_right):
        cell = tblNew.cell(i, 1)
        cell.paragraphs[0].text = text


#ADD ADDRESS 
def address_add_additional(doc):
    doc.add_paragraph()
    type_of_change_heading = doc.add_paragraph("TYPE OF CHANGE: ADDING AN ADDITIONAL (NEW) ADDRESS")
    type_of_change_heading.style = doc.styles["Heading 2"]

    tblNew = doc.add_table(rows=20, cols=2)
    tblNew.style = "Table Grid"

    cell_texts_left = [
    "ADDITIONAL (NEW) ADDRESS DETAILS",
    "Address Line 1:",
    "Address Line 2:",
    "Address Line 3:",
    "City:",
    "County:",
    "Province:",
    "State:",
    "Postal Code:",
    "Country:",
    "ADDRESS PURPOSE (Check address purpose below)",
    "__ PURCHASING ADDRESS",
    "__ PAY ADDRESS",
    "Definition for 'PURCHASING ADDRESS': For Indirect Suppliers it is the PURCHASE ORDERS MAILING ADDRESS, For Direct suppliers it is the SHIP FROM ADDRESS.",
    "",
    "IF ADDING AN ADDITIONAL PURCHASING ADDRESS, SUPPLIER CONTACT TO RECEIVE PURCHASE ORDERS (POs)",
    "First/Last Name:",
    "Email Address:",
    "Phone Number:",
    "Fax Number (If applicable):"
    ]


    for i, text in enumerate(cell_texts_left):
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text

    cell_texts_right  = [
    "IF ADDING AN ADDITIONAL PURCHASING ADDRESS, THE ASSOCIATED PAY ADDRESS MUST BE PROVIDED",
    "Address Line 1:",
    "Address Line 2:",
    "Address Line 3:",
    "City:",
    "County:",
    "Province:",
    "State:",
    "Postal Code:",
    "Country:",
    "",
    "",
    "",
    "Definition for 'PAY ADDRESS': Remit to address that appears on a supplier's invoice.",
    "",
    "IF ADDING AN ADDITIONAL PAY ADDRESS, SUPPLIER CONTACT TO RECEIVE PAYMENTS",
    "First/Last Name:",
    "Email Address:",
    "Phone Number:",
    "Fax Number (If applicable):"
    ]

    for i, text in enumerate(cell_texts_right):
        cell = tblNew.cell(i, 1)
        cell.paragraphs[0].text = text

#CURRENCY UPDATE
def currency_update(doc):
    print("CURRENCY")
    doc.add_paragraph()
    type_of_change_heading = doc.add_paragraph("TYPE OF CHANGE:  ENTIRE SUPPLIER CURRENCY CHANGE ")
    type_of_change_heading.style = doc.styles["Heading 2"]
    doc.add_paragraph("(Note: The currency of an existing site cannot be changed for any region.  A new site for the affected address(es) must be created with the new currency.) ")


    tblNew = doc.add_table(rows=2, cols=2)
    tblNew.style = "Table Grid"

    cell_texts_left = [
        "PREVIOUS CURRENCY"
    ]


    for i, text in enumerate(cell_texts_left):
        cell = tblNew.cell(i, 0)
        cell.paragraphs[0].text = text

    cell_texts_right  = [
         "NEW CURRENCY"
    ]

    for i, text in enumerate(cell_texts_right):
        cell = tblNew.cell(i, 1)
        cell.paragraphs[0].text = text


