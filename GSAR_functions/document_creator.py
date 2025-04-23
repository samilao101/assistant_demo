
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from GSAR_functions import GSAR_template as template
from GSAR_functions import GSAR_eForm_template as eform_template


def create_documents(settings):
    print("Creating documents...")

    # Create a new Document
    # Create a new Document
    doc = Document()

    # Add a title

    template.create_headings(doc)
    template.create_ustax_id(doc)

    # Add paragraphs
    if settings.get("name_update") == "yes":
        template.name_update(doc)
    if settings.get("address_update") == "yes":
        template.address_replace_update(doc)
    if settings.get("address_add") == "yes":
        template.address_add_additional(doc)
    if settings.get("banking_update") == "yes":
        template.banking_update(doc)
    if settings.get("contact_update") == "yes":
        template.contact_replace_update(doc)
    if settings.get("contact_add") == "yes":
        template.contact_add_additional(doc)
    if settings.get("currency_update") == "yes":
        print("cureency")
        template.currency_update(doc)

    template.create_signature(doc)

    doc.save('generated_document.docx')

def return_document():
    doc = Document("generated_document.docx")
    return doc 

def create_eform_documents(settings):
    print("Creating documents...")

    # Create a new Document
    # Create a new Document
    doc = Document()

    # Add a title

    eform_template.create_headings(doc)
    eform_template.create_ustax_id(doc)

    # Add paragraphs
    if settings.get("name_update") == "yes":
        eform_template.name_update(doc)
    if settings.get("address_update") == "yes":
        eform_template.address_replace_update(doc)
    if settings.get("address_add") == "yes":
        eform_template.address_add_additional(doc)
    if settings.get("banking_update") == "yes":
        eform_template.banking_update(doc)
    if settings.get("contact_update") == "yes":
        eform_template.contact_replace_update(doc)
    if settings.get("contact_add") == "yes":
        eform_template.contact_add_additional(doc)
    if settings.get("currency_update") == "yes":
        print("cureency")
        eform_template.currency_update(doc)

    eform_template.create_signature(doc)

    doc.save('generated_eform_document.docx')



def return_eform_document():
    doc = Document("generated_eform_document.docx")
    return doc 

